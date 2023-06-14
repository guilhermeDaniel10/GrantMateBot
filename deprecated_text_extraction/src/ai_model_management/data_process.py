
import nltk
import collections
import spacy
import fitz
import os
import string
import pandas as pd
import re
import numpy as np
import json
import pdfquery
import PyPDF2
import docx
import requests


from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from unidecode import unidecode
from dotenv import load_dotenv
from operator import itemgetter
from grobid_client import Client
from grobid_client.api.pdf import process_fulltext_document
from grobid_client.models import Article, ProcessForm
from grobid_client.types import TEI, File
import urllib3

from ai_model_management.document_blocks import DocumentBlock

FILE_STORAGE_PATH_PYTHON = os.getenv("FILE_STORAGE_PATH_PYTHON")

load_dotenv()

# This class turned out to be a pdf file data handling, since all the files must be in pdf to extract the data


class DataProcess:
    def __init__(self):
        nltk.download('punkt')
        nltk.download('averaged_perceptron_tagger')

    # uses NER to find relevant words in a document
    def extract_text_keywords(self, content):
        # load spacy model
        nlp = spacy.load('en_core_web_lg')
        doc = nlp(content)

        for ent in doc.ents:
            print(ent.text, ent.start_char, ent.end_char, ent.label_)

    # extracts the most common words of a document
    def extract_most_used_words(self, content, language, n_words):
        text = content.lower()
        text = text.translate(str.maketrans('', '', string.punctuation))
        tokens = word_tokenize(text)

        stop_words = set(stopwords.words(language))
        filtered_tokens = [word for word in tokens if word not in stop_words]

        word_freq = collections.Counter(filtered_tokens)

        most_common_words = word_freq.most_common(n_words)

        # Return a list of the most common words
        return [word for word, frequency in most_common_words]

    # UGLY FUNCTION- TODO: MAKE IT SIMPLER

    def document_block_extraction(self, filename):

        doc = fitz.open(FILE_STORAGE_PATH_PYTHON + filename)
        block_dict = {}
        page_num = 1

        for page in doc:  # Iterate all pages in the document
            file_dict = page.get_text('dict')  # Get the page dictionary
            block = file_dict['blocks']  # Get the block information
            block_dict[page_num] = block  # Store in block dictionary
            page_num += 1  # Increase the page value by 1

        spans = pd.DataFrame(
            columns=['xmin', 'ymin', 'xmax', 'ymax', 'text', 'tag'])

        print(spans, flush=True)
        rows = []

        for page_num, blocks in block_dict.items():
            for block in blocks:
                if block['type'] == 0:
                    for line in block['lines']:
                        for span in line['spans']:
                            xmin, ymin, xmax, ymax = list(span['bbox'])
                            font_size = span['size']
                            text = unidecode(span['text'])
                            span_font = span['font']
                            is_upper = False
                            is_bold = False
                            if "bold" in span_font.lower():
                                is_bold = True

                            if re.sub("[\(\[].*?[\)\]]", "", text).isupper():
                                is_upper = True

                            if text.replace(" ", "") != "":
                                rows.append(
                                    (xmin, ymin, xmax, ymax, text, is_upper, is_bold, span_font, font_size))

        span_df = pd.DataFrame(rows, columns=[
                               'xmin', 'ymin', 'xmax', 'ymax', 'text', 'is_upper', 'is_bold', 'span_font', 'font_size'])
        print(span_df, flush=True)

        span_df.to_csv(FILE_STORAGE_PATH_PYTHON + "first_span.csv",
                       encoding='utf-8', index=True)

        span_scores = []
        span_num_occur = {}
        special = '[(_:/,#%\=@)]'

        for index, span_row in span_df.iterrows():
            score = round(span_row.font_size)
            text = span_row.text

            if not re.search(special, text):
                if span_row.is_bold:
                    score += 1
                if span_row.is_upper:
                    score += 1

            span_scores.append(score)

        values, counts = np.unique(span_scores, return_counts=True)

        style_dict = {}
        for value, count in zip(values, counts):
            style_dict[value] = count
        sorted(style_dict.items(), key=lambda x: x[1])
        p_size = max(style_dict, key=style_dict.get)
        idx = 0
        tag = {}

        for size in sorted(values, reverse=True):

            idx += 1
            if size == p_size:
                idx = 0
                tag[size] = 'p'
            if size > p_size:
                tag[size] = 'h{0}'.format(idx)
            if size < p_size:
                tag[size] = 's{0}'.format(idx)

        span_tags = [tag[score] for score in span_scores]
        span_df['tag'] = span_tags

        print(span_df, flush=True)
        span_df.to_csv(FILE_STORAGE_PATH_PYTHON + "spans.csv",
                       encoding='utf-8', index=True)

        span_scores = []
        span_num_occur = {}
        special = '[(_:/,#%\=@)]'

        for index, span_row in span_df.iterrows():
            score = round(span_row.font_size)
            text = span_row.text

            if not re.search(special, text):
                if span_row.is_bold:
                    score += 1
                if span_row.is_upper:
                    score += 1

            span_scores.append(score)
        values, counts = np.unique(span_scores, return_counts=True)

        style_dict = {}
        for value, count in zip(values, counts):
            style_dict[value] = count
        sorted(style_dict.items(), key=lambda x: x[1])

        p_size = max(style_dict, key=style_dict.get)
        idx = 0
        tag = {}

        for size in sorted(values, reverse=True):
            idx += 1
            if size == p_size:
                idx = 0
                tag[size] = 'p'
            if size > p_size:
                tag[size] = 'h{0}'.format(idx)
            if size < p_size:
                tag[size] = 's{0}'.format(idx)

        span_tags = [tag[score] for score in span_scores]
        span_df['tag'] = span_tags

        print(span_df, flush=True)

        headings_list = []
        text_list = []
        tmp = []
        heading = ''

        for index, span_row in span_df.iterrows():
            text = span_row.text
            tag = span_row.tag

            if 'h' in tag:
                headings_list.append(text)
                text_list.append('\n'.join(tmp))
                tmp = []
                heading = text
            else:
                tmp.append(text)

        text_list.append('\n'.join(tmp))
        text_list = text_list[1:]
        text_df = pd.DataFrame(zip(headings_list, text_list), columns=[
                               'heading', 'content'])

        print(text_df)

        text_df.to_csv(FILE_STORAGE_PATH_PYTHON + "result.csv",
                       encoding='utf-8', index=True)

    def extract_to_html(self, filename):
        doc = fitz.open(FILE_STORAGE_PATH_PYTHON + filename)
        for page in doc:  # Iterate all pages in the document
            for line in page.get_text("html").splitlines():
                print(line, flush=True)

    def extract_in_block_order(self, filename):
        # Load the English language model
        nlp = spacy.load('en_core_web_sm')

        # Open the PDF file in read-binary mode
        with open(FILE_STORAGE_PATH_PYTHON + filename, 'rb') as pdf_file:
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Initialize the dataframe
            df = pd.DataFrame(columns=['heading', 'paragraph'])

            # Loop over the pages in the PDF file
            for page_num in range(len(pdf_reader.pages)):
                # Get the page object
                page = pdf_reader.pages[page_num]

                # Extract the text from the page
                text = page.extract_text()

                # Use spaCy to parse the text
                doc = nlp(text)

                # Loop over the sentences in the parsed text
                for sent in doc.sents:
                    # Check if the sentence is a heading
                    if sent.text.isupper() or re.match(r'^\d+(\.\d+)*\s+', sent.text):
                        heading = sent.text

                        # Initialize the paragraph
                        paragraph = ''

                        # Loop over the following sentences until the next heading
                        for next_sent in doc[sent.end:].sents:
                            if next_sent.text.isupper() or re.match(r'^\d+(\.\d+)*\s+', next_sent.text):
                                break
                            paragraph += next_sent.text

                        # Append the heading and paragraph to the dataframe
                        # Check if the paragraph is non-empty before appending
                        if paragraph.strip():
                            df.loc[len(df)] = {
                                'heading': heading, 'paragraph': paragraph}
                        paragraph = ''
                        # df = df.append({'heading': heading, 'paragraph': paragraph}, ignore_index=True)

        df.to_csv(FILE_STORAGE_PATH_PYTHON + "result.csv",
                  encoding='utf-8', index=True)
        # Print the dataframe
        print(df)

    def extract_text_from_pdf(self, filename):
        # Open the PDF file
        with fitz.open(FILE_STORAGE_PATH_PYTHON + filename) as doc:
            # Initialize the dataframe
            df = pd.DataFrame(columns=['heading', 'paragraph'])

            # Loop over the pages in the PDF file
            for page_num in range(doc.page_count):
                # Get the page object
                page = doc.load_page(page_num)

                # Extract the text from the page
                text = page.get_text()

                # Split the text into sentences
                sentences = text.split('\n')

                # Loop over the sentences in the text
                for i in range(len(sentences)):
                    # Check if the sentence is a heading
                    if sentences[i].isupper() or re.match(r'^\d+(\.\d+)*\s+', sentences[i]):
                        heading = sentences[i]

                        # Initialize the paragraph
                        paragraph = ''

                        # Loop over the following sentences until the next heading
                        j = i + 1
                        while j < len(sentences) and not (sentences[j].isupper() or re.match(r'^\d+(\.\d+)*\s+', sentences[j])):
                            paragraph += sentences[j] + ' '
                            j += 1

                        # Append the heading and paragraph to the dataframe
                        df = pd.concat([df, pd.DataFrame.from_records(
                            [{'heading': heading, 'paragraph': paragraph}])])

        # Reset the index of the dataframe
        df.reset_index(drop=True, inplace=True)

        # Print the dataframe
        print(df)

        # Write the dataframe to a CSV file
        df.to_csv(FILE_STORAGE_PATH_PYTHON + 'output.csv', index=False)

    def extract_pdf_structured(self, filename):
        # Open the PDF file
        with fitz.open(FILE_STORAGE_PATH_PYTHON + filename) as doc:
            # Initialize the dataframe
            df = pd.DataFrame(columns=['heading', 'paragraph'])

            for page_num in range(doc.page_count):
                # Get the page object
                page = doc.load_page(page_num)

                # Extract the text blocks from the page
                blocks = page.get_text("blocks")

                # Loop over the blocks in the text
                for block in blocks:
                    print(block, flush=True)
                    # Check if the block is a heading
                    if block['type'] == 0 and (block['text'].isupper() or re.match(r'^\d+(\.\d+)*\s+', block['text'])):
                        heading = block['text']

                        # Initialize the paragraph
                        paragraph = ''

                        # Loop over the following blocks until the next heading
                        for i in range(block['index'] + 1, len(blocks)):
                            if blocks[i]['type'] == 0 and (blocks[i]['text'].isupper() or re.match(r'^\d+(\.\d+)*\s+', blocks[i]['text'])):
                                break
                            paragraph += blocks[i]['text'] + ' '

                        # Append the heading and paragraph to the dataframe
                        df = pd.concat([df, pd.DataFrame.from_records(
                            [{'heading': heading, 'paragraph': paragraph}])])

        # Reset the index of the dataframe
        df.reset_index(drop=True, inplace=True)

        # Print the dataframe
        print(df)

        # Write the dataframe to a CSV file
        df.to_csv('output.csv', index=False)

    def extract_document_to_dataframe(self, filename):  # has to be docx
        # Open the Word document
        doc = docx.Document(FILE_STORAGE_PATH_PYTHON + filename)

        current_heading = ''
        current_paragraph = ''
        blocks = []

        is_heading = False

        for paragraph in doc.paragraphs:
            runboldtext = ''
            if "Heading" in paragraph.style.name:
                is_heading = True
            else:
                for run in paragraph.runs:
                    if run.bold:
                        runboldtext = runboldtext + run.text

                # Now check if the value of "runboldtext" matches the entire paragraph text. If it matches, it means all the words in the current paragraph are bold and can be considered as a heading
                if runboldtext == str(paragraph.text) and runboldtext != '':
                    is_heading = True
                else:
                    is_heading = False

            if is_heading == True:
                document_block = DocumentBlock(
                    current_heading, current_paragraph)
                blocks.append(document_block)
                current_heading = str(paragraph.text)
                current_paragraph = ''
            else:
                current_paragraph += str(paragraph.text)

        df = pd.DataFrame.from_records([block.to_dict() for block in blocks])

        # Write the dataframe to a CSV file
        df.to_csv(FILE_STORAGE_PATH_PYTHON + 'output.csv', sep="|", index=True)

        print(df, flush=True)
        return df

    def extract_word_to_dataframe(self, filename):
        # Open the Word document
        doc = docx.Document(FILE_STORAGE_PATH_PYTHON + filename)

        # Initialize variables to store the extracted information
        current_heading = None
        current_subheading = None
        current_paragraphs = []

        # Initialize empty lists to store the extracted information
        headings = []
        paragraphs = []

        paragraph_content = ""

        # Iterate through each paragraph in the document
        for paragraph in doc.paragraphs:

            # Perform the below logic only for paragraph content which does not have its native style as "Heading"
            if "Heading" not in paragraph.style.name:
                # Start by initializing an empty string to store bold words inside a run
                runboldtext = ''

                # Iterate over all runs of the current paragraph and collect all the words which are bold into the variable "runboldtext"
                for run in paragraph.runs:
                    if run.bold:
                        runboldtext = runboldtext + run.text

                # Now check if the value of "runboldtext" matches the entire paragraph text. If it matches, it means all the words in the current paragraph are bold and can be considered as a heading
                if runboldtext == str(paragraph.text) and runboldtext != '':
                    current_heading = str(paragraph.text)
                    current_paragraphs = []
                else:
                    current_paragraphs.append(str(paragraph.text))
            else:
                current_heading = str(paragraph.text)
                current_paragraphs = []

            # Add the current heading and paragraph content to the respective lists
            if current_heading is not None and current_paragraphs:
                # Check if the current heading already exists in the headings list
                if current_heading in headings:
                    # If the heading already exists, append the paragraph content to the corresponding row in the dataframe
                    idx = headings.index(current_heading)
                    paragraphs[idx] += ' ' + ' '.join(current_paragraphs)
                else:
                    # If the heading is new, add it and its corresponding paragraph content to the lists
                    headings.append(current_heading)
                    paragraphs.append(' '.join(current_paragraphs))

        # Create a pandas dataframe with the extracted information
        df = pd.DataFrame(
            {'heading': headings, 'paragraph_content': paragraphs})

        # Write the dataframe to a CSV file
        df.to_csv(FILE_STORAGE_PATH_PYTHON + 'output.csv', index=False)
        return df

    def extract_data_via_dataframe(self, filename):
        # First, create an opener which accepts a PDF file path
        opener = open(FILE_STORAGE_PATH_PYTHON + filename, 'rb')
        # Second, read the opened file
        pdf_file_reader = PyPDF2.PdfReader(opener)

        def pdf_info(read_pdf):
            pdf_info_dict = {}
            pdf_info = {}
            for key, value in read_pdf.metadata.items():
                pdf_info_dict[re.sub('/', "", key)] = value
                return pdf_info_dict

        def pdf_list_to_series_and_df(pdf_info_dict):
            pdf_series = pd.Series(pdf_info_dict)

            key_list = []
            val_list = []
            for key, val in pdf_info_dict.items():
                key_list.append(key)
                val_list.append(val)
            pdf_df = pd.Series.to_frame(pdf_series)
            pdf_df = pd.DataFrame(
                {"Attribute": key_list, "Informaton": val_list}, index=key_list)
            return pdf_series, pdf_df

        read_pdf = pdf_info(pdf_file_reader)
        series, df = pdf_list_to_series_and_df(read_pdf)

        print(series, flush=True)

    def extract_text_using_nlp(self, filename):
        nlp = spacy.load('en_core_web_sm')
        pdf_file = open(FILE_STORAGE_PATH_PYTHON + filename, 'rb')
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        text = ''
        for page in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page].extract_text()

        section_titles = []
        for token in nlp(text):
            if token.is_title:
                section_titles.append(token.text)
                print(token, flush=True)

        file_path = open(FILE_STORAGE_PATH_PYTHON + 'items.txt', 'w')
        for item in section_titles:
            file_path.write(item+"\n")
        file_path.close()

    def extract_using_grobid(self, filename):
        '''
        grobid_url = "http://grobid:8070/api/processFulltextDocument"

        payload = {'input': FILE_STORAGE_PATH_PYTHON + filename, 'consolidateHeader': '1'}
        headers = {}

        response = requests.request("GET", grobid_url, headers=headers, params=payload)

        print(response.text)
        '''
        
        '''
        client = Client(base_url="http://grobid:8070/api")
        #pdf_file = open(FILE_STORAGE_PATH_PYTHON + filename, 'rb')
        with open(FILE_STORAGE_PATH_PYTHON + filename, 'rb') as fin:
            form = ProcessForm(
                segment_sentences="1",
                input_=File(file_name=filename, payload=fin,
                            mime_type="application/pdf"),
            )
            r = process_fulltext_document.sync_detailed(
                client=client, multipart_data=form, timeout=30)
            if r.is_success:
                article: Article = TEI.parse(r.content, figures=False)
                assert article.title
        '''

        url = 'http://grobid:8070/api/processFulltextDocument'
        filename_without_extension = os.path.splitext(filename)[0]
        filename = FILE_STORAGE_PATH_PYTHON + filename
        

        # Open the file in binary mode and read its contents
        with open(filename, 'rb') as f:
            pdf_content = f.read()

        # Define the payload as a dictionary with the necessary parameters
        payload = {
            'input': ('file.pdf', pdf_content, 'application/pdf'),
            'fulltext': '1',
            'consolidateHeader': '1'
        }

        # Make the POST request and capture the response
        response = requests.post(url, files=payload)

        # Check the status code of the response
        if response.status_code == 200:
            # Extract the TEI-XML content from the response
            tei_xml = response.content
            print(tei_xml)

            with open(FILE_STORAGE_PATH_PYTHON + filename_without_extension + ".xml", "wb") as f:
                f.write(tei_xml)
        else:
            print(f'Error: {response.status_code}')

    
    def export_pdf_into_text_blocks(self, filename):
         with fitz.open(FILE_STORAGE_PATH_PYTHON + filename) as doc:
            iterator = 0
            for page_num in range(doc.page_count):
                # Get the page object
                page = doc.load_page(page_num)

                # Extract the text blocks from the page
                blocks = page.get_text("blocks")

                # Loop over the blocks in the text
                for block in blocks:
                    print(block, flush=True)
                    print(iterator, flush= True)
